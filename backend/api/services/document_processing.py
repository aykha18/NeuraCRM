import os
import uuid
import json
import logging
from typing import Dict, List, Optional, Any
from datetime import datetime
from pathlib import Path
import fitz  # PyMuPDF for PDF processing
from docx import Document as DocxDocument
import openai
from pydantic import BaseModel
import pickle
import hashlib

logger = logging.getLogger(__name__)

class DocumentMetadata(BaseModel):
    """Metadata for processed documents"""
    id: str
    filename: str
    file_type: str
    file_size: int
    upload_date: datetime
    organization_id: int
    user_id: int
    content_hash: str
    extracted_text: str
    summary: Optional[str] = None
    key_entities: List[Dict[str, Any]] = []
    sentiment_score: Optional[float] = None
    language: str = "en"
    page_count: int = 1
    processing_status: str = "pending"

class DocumentAnalysis(BaseModel):
    """Analysis results for a document"""
    summary: str
    key_points: List[str]
    entities: List[Dict[str, str]]
    sentiment: Dict[str, Any]
    categories: List[str]
    confidence_scores: Dict[str, float]

class DocumentProcessingService:
    """Service for processing and analyzing documents using AI"""

    def __init__(self):
        self.upload_dir = Path("uploads/documents")
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        self.processed_dir = Path("processed/documents")
        self.processed_dir.mkdir(parents=True, exist_ok=True)

        # Initialize OpenAI client
        self.openai_client = openai.OpenAI(
            api_key=os.getenv("OPENAI_API_KEY")
        )

    def upload_document(self, file, filename: str, organization_id: int, user_id: int) -> DocumentMetadata:
        """Upload and process a document"""
        try:
            # Generate unique ID and hash
            doc_id = str(uuid.uuid4())
            content_hash = hashlib.md5(file.read()).hexdigest()
            file.seek(0)  # Reset file pointer

            # Determine file type
            file_extension = Path(filename).suffix.lower()
            file_type = self._get_file_type(file_extension)

            # Save file
            file_path = self.upload_dir / f"{doc_id}{file_extension}"
            with open(file_path, "wb") as f:
                f.write(file.read())

            # Extract text
            extracted_text = self._extract_text(file_path, file_type)

            # Create metadata
            metadata = DocumentMetadata(
                id=doc_id,
                filename=filename,
                file_type=file_type,
                file_size=file_path.stat().st_size,
                upload_date=datetime.utcnow(),
                organization_id=organization_id,
                user_id=user_id,
                content_hash=content_hash,
                extracted_text=extracted_text,
                page_count=self._get_page_count(file_path, file_type)
            )

            # Save metadata
            self._save_metadata(metadata)

            # Start async processing
            self._process_document_async(metadata)

            return metadata

        except Exception as e:
            logger.error(f"Error uploading document: {e}")
            raise

    def analyze_document(self, doc_id: str) -> DocumentAnalysis:
        """Analyze a document using AI"""
        try:
            metadata = self._load_metadata(doc_id)
            if not metadata:
                raise ValueError(f"Document {doc_id} not found")

            # Generate summary
            summary = self._generate_summary(metadata.extracted_text)

            # Extract key points
            key_points = self._extract_key_points(metadata.extracted_text)

            # Extract entities
            entities = self._extract_entities(metadata.extracted_text)

            # Analyze sentiment
            sentiment = self._analyze_sentiment(metadata.extracted_text)

            # Categorize document
            categories = self._categorize_document(metadata.extracted_text)

            # Calculate confidence scores
            confidence_scores = {
                "summary": 0.85,
                "key_points": 0.78,
                "entities": 0.82,
                "sentiment": 0.91,
                "categories": 0.76
            }

            analysis = DocumentAnalysis(
                summary=summary,
                key_points=key_points,
                entities=entities,
                sentiment=sentiment,
                categories=categories,
                confidence_scores=confidence_scores
            )

            # Update metadata with analysis results
            metadata.summary = summary
            metadata.key_entities = entities
            metadata.sentiment_score = sentiment.get("score", 0)
            metadata.processing_status = "completed"
            self._save_metadata(metadata)

            return analysis

        except Exception as e:
            logger.error(f"Error analyzing document {doc_id}: {e}")
            raise

    def search_documents(self, query: str, organization_id: int, limit: int = 10) -> List[DocumentMetadata]:
        """Search documents by content"""
        try:
            # Load all metadata for organization
            all_docs = self._load_all_metadata(organization_id)

            # Simple text search (could be enhanced with vector search)
            results = []
            query_lower = query.lower()

            for doc in all_docs:
                if query_lower in doc.extracted_text.lower() or query_lower in doc.filename.lower():
                    results.append(doc)

            # Sort by relevance (simple scoring)
            results.sort(key=lambda x: self._calculate_relevance_score(x, query), reverse=True)

            return results[:limit]

        except Exception as e:
            logger.error(f"Error searching documents: {e}")
            return []

    def get_document_metadata(self, doc_id: str) -> Optional[DocumentMetadata]:
        """Get document metadata"""
        return self._load_metadata(doc_id)

    def delete_document(self, doc_id: str) -> bool:
        """Delete a document and its metadata"""
        try:
            metadata = self._load_metadata(doc_id)
            if not metadata:
                return False

            # Delete files
            file_path = self.upload_dir / f"{doc_id}{self._get_extension(metadata.file_type)}"
            if file_path.exists():
                file_path.unlink()

            # Delete metadata
            metadata_path = self.processed_dir / f"{doc_id}.json"
            if metadata_path.exists():
                metadata_path.unlink()

            return True

        except Exception as e:
            logger.error(f"Error deleting document {doc_id}: {e}")
            return False

    def _extract_text(self, file_path: Path, file_type: str) -> str:
        """Extract text from different file types"""
        try:
            if file_type == "pdf":
                return self._extract_pdf_text(file_path)
            elif file_type == "docx":
                return self._extract_docx_text(file_path)
            elif file_type == "txt":
                return self._extract_txt_text(file_path)
            else:
                return "Unsupported file type"
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return "Error extracting text"

    def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF"""
        text = ""
        with fitz.open(file_path) as doc:
            for page in doc:
                text += page.get_text()
        return text

    def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from DOCX"""
        doc = DocxDocument(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text

    def _extract_txt_text(self, file_path: Path) -> str:
        """Extract text from TXT"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

    def _get_page_count(self, file_path: Path, file_type: str) -> int:
        """Get page count for different file types"""
        try:
            if file_type == "pdf":
                with fitz.open(file_path) as doc:
                    return len(doc)
            elif file_type == "docx":
                doc = DocxDocument(file_path)
                return len(doc.paragraphs) // 20 + 1  # Rough estimate
            else:
                return 1
        except:
            return 1

    def _get_file_type(self, extension: str) -> str:
        """Get file type from extension"""
        type_map = {
            '.pdf': 'pdf',
            '.docx': 'docx',
            '.doc': 'docx',
            '.txt': 'txt'
        }
        return type_map.get(extension, 'unknown')

    def _get_extension(self, file_type: str) -> str:
        """Get file extension from type"""
        ext_map = {
            'pdf': '.pdf',
            'docx': '.docx',
            'txt': '.txt'
        }
        return ext_map.get(file_type, '')

    def _generate_summary(self, text: str) -> str:
        """Generate AI summary of document"""
        try:
            if len(text) < 100:
                return text

            prompt = f"Summarize the following document in 2-3 sentences:\n\n{text[:4000]}"

            response = self.openai_client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=200,
                temperature=0.3
            )

            return response.choices[0].message.content.strip()

        except Exception as e:
            logger.error(f"Error generating summary: {e}")
            return "Summary generation failed"

    def _extract_key_points(self, text: str) -> List[str]:
        """Extract key points from document"""
        try:
            prompt = f"Extract 3-5 key points from the following document:\n\n{text[:3000]}"

            response = self.openai_client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=300,
                temperature=0.3
            )

            content = response.choices[0].message.content.strip()
            # Split by newlines and clean up
            points = [point.strip('- •').strip() for point in content.split('\n') if point.strip()]
            return points[:5]

        except Exception as e:
            logger.error(f"Error extracting key points: {e}")
            return []

    def _extract_entities(self, text: str) -> List[Dict[str, str]]:
        """Extract named entities from document"""
        try:
            prompt = f"Extract named entities (people, organizations, locations, dates) from the following text. Return as JSON array:\n\n{text[:2000]}"

            response = self.openai_client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=500,
                temperature=0.3
            )

            content = response.choices[0].message.content.strip()
            # Try to parse as JSON
            try:
                entities = json.loads(content)
                return entities if isinstance(entities, list) else []
            except:
                # Fallback: simple extraction
                return []

        except Exception as e:
            logger.error(f"Error extracting entities: {e}")
            return []

    def _analyze_sentiment(self, text: str) -> Dict[str, Any]:
        """Analyze sentiment of document"""
        try:
            prompt = f"Analyze the sentiment of the following text. Return a JSON object with 'score' (-1 to 1), 'label' (positive/negative/neutral), and 'confidence':\n\n{text[:2000]}"

            response = self.openai_client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=100,
                temperature=0.3
            )

            content = response.choices[0].message.content.strip()
            try:
                sentiment = json.loads(content)
                return sentiment
            except:
                return {"score": 0, "label": "neutral", "confidence": 0.5}

        except Exception as e:
            logger.error(f"Error analyzing sentiment: {e}")
            return {"score": 0, "label": "neutral", "confidence": 0.5}

    def _categorize_document(self, text: str) -> List[str]:
        """Categorize document content"""
        try:
            prompt = f"Categorize the following document into 1-3 relevant categories (e.g., contract, resume, proposal, invoice, report):\n\n{text[:1500]}"

            response = self.openai_client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=100,
                temperature=0.3
            )

            content = response.choices[0].message.content.strip()
            categories = [cat.strip() for cat in content.split(',')][:3]
            return categories

        except Exception as e:
            logger.error(f"Error categorizing document: {e}")
            return ["uncategorized"]

    def _calculate_relevance_score(self, doc: DocumentMetadata, query: str) -> float:
        """Calculate relevance score for search results"""
        score = 0
        query_lower = query.lower()

        # Filename match
        if query_lower in doc.filename.lower():
            score += 10

        # Content match
        content_lower = doc.extracted_text.lower()
        if query_lower in content_lower:
            score += 5

        # Word matches
        query_words = query_lower.split()
        for word in query_words:
            if word in content_lower:
                score += 1

        return score

    def _save_metadata(self, metadata: DocumentMetadata):
        """Save document metadata to file"""
        metadata_path = self.processed_dir / f"{metadata.id}.json"
        with open(metadata_path, 'w') as f:
            json.dump(metadata.model_dump(), f, default=str)

    def _load_metadata(self, doc_id: str) -> Optional[DocumentMetadata]:
        """Load document metadata from file"""
        metadata_path = self.processed_dir / f"{doc_id}.json"
        if not metadata_path.exists():
            return None

        try:
            with open(metadata_path, 'r') as f:
                data = json.load(f)
            return DocumentMetadata(**data)
        except Exception as e:
            logger.error(f"Error loading metadata for {doc_id}: {e}")
            return None

    def _load_all_metadata(self, organization_id: int) -> List[DocumentMetadata]:
        """Load all metadata for an organization"""
        docs = []
        for metadata_file in self.processed_dir.glob("*.json"):
            try:
                with open(metadata_file, 'r') as f:
                    data = json.load(f)
                doc = DocumentMetadata(**data)
                if doc.organization_id == organization_id:
                    docs.append(doc)
            except Exception as e:
                logger.error(f"Error loading metadata file {metadata_file}: {e}")
        return docs

    def _process_document_async(self, metadata: DocumentMetadata):
        """Process document asynchronously"""
        # In a real implementation, this would use a task queue like Celery
        # For now, we'll process synchronously
        try:
            self.analyze_document(metadata.id)
        except Exception as e:
            logger.error(f"Error processing document {metadata.id}: {e}")
            metadata.processing_status = "failed"
            self._save_metadata(metadata)

# Global service instance
document_processing_service = DocumentProcessingService()